"""Generate MVR_TEMPLATE.docx with {{TOKEN}} placeholders for JSZip fill.
   Visual styling matched to CLAIM_SUMMARY_TEMPLATE.docx (navy/gold scheme).
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand colors (from Claim Summary template) ──
NAVY_FILL  = '1B2A4A'   # section heading bars (matched from Claim Summary)
NAVY_TEXT  = '0D1B2A'   # title text
GOLD_TEXT  = 'C9A84C'   # gold accent
WHITE      = 'FFFFFF'
GRAY_BODY  = '333333'
GRAY_SUB   = '888888'
GRAY_FOOT  = '999999'
FILL_LABEL     = 'F5F6F8'
FILL_HIGHLIGHT = 'FFF8E7'
FILL_GOLD_TINT = 'E8D5A3'
BORDER_CELL    = 'CCCCCC'
BORDER_GOLD    = 'C8A951'
BORDER_FOOTER  = 'E8EAED'
FONT = 'Arial'

doc = Document()

section = doc.sections[0]
section.top_margin    = Inches(0.5)
section.bottom_margin = Inches(0.55)
section.left_margin   = Inches(0.6)
section.right_margin  = Inches(0.6)

# ── Helpers ──

def set_run_font(run, name=FONT):
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rFonts.set(qn(attr), name)

def make_run(para, text, size=9, bold=False, color=GRAY_BODY, italic=False):
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(int(color[:2],16), int(color[2:4],16), int(color[4:],16))
    set_run_font(run)
    return run

def set_shd(element, fill):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    element.append(shd)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    set_shd(tcPr, fill)

def set_cell_borders(cell, color=BORDER_CELL, sz='1'):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:color'), color)
        el.set(qn('w:sz'), sz)
        borders.append(el)
    tcPr.append(borders)

def set_cell_text(cell, text, bold=False, size=9, color=GRAY_BODY):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    make_run(p, text, size=size, bold=bold, color=color)

def styled_cell(cell, text, bold=False, size=9, color=GRAY_BODY, fill=None, align=None):
    set_cell_text(cell, text, bold=bold, size=size, color=color)
    if fill:
        set_cell_shading(cell, fill)
    set_cell_borders(cell)
    if align:
        cell.paragraphs[0].alignment = align

def section_heading(doc, text):
    """Navy banner — #1B2A4A bg, white bold text (matches Claim Summary)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    set_shd(pPr, NAVY_FILL)
    make_run(p, text, size=9, bold=True, color=WHITE)
    return p

def add_para_border(para, edge='bottom', color=BORDER_GOLD, sz='4', space='4'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el = OxmlElement(f'w:{edge}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:color'), color)
    el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), space)
    pBdr.append(el)
    pPr.append(pBdr)

def add_right_tab(para, pos='9026'):
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), pos)
    tabs.append(tab)
    pPr.append(tabs)

def set_spacing(para, before=None, after=None):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    if before is not None:
        sp.set(qn('w:before'), str(before))
    if after is not None:
        sp.set(qn('w:after'), str(after))
    pPr.append(sp)

# =============================================
#  PAGE HEADER — {{HEADER_BUSINESS}} | Powered by Claim Cipher™
# =============================================

header = section.header
header.is_linked_to_previous = False
hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
hp.clear()

add_para_border(hp, edge='bottom', color=BORDER_GOLD, sz='4', space='4')
set_spacing(hp, after=0)
add_right_tab(hp)

make_run(hp, '{{HEADER_BUSINESS}}', size=8, bold=True, color=NAVY_FILL)
make_run(hp, '   |   ', size=7, color=GRAY_SUB)
make_run(hp, 'Powered by Claim Cipher\u2122', size=7, color=GOLD_TEXT)

# =============================================
#  PAGE FOOTER — {{HEADER_BUSINESS}} • Powered by Claim Cipher™  Page X
# =============================================

footer_sec = section.footer
footer_sec.is_linked_to_previous = False
fp = footer_sec.paragraphs[0] if footer_sec.paragraphs else footer_sec.add_paragraph()
fp.clear()

add_para_border(fp, edge='top', color=BORDER_FOOTER, sz='2', space='4')
set_spacing(fp, before=0)
add_right_tab(fp)

make_run(fp, '{{HEADER_BUSINESS}}  \u2022  Powered by Claim Cipher\u2122', size=7, color=GRAY_FOOT)

# Tab to right side for page number
tab_run = fp.add_run('\tPage ')
tab_run.font.size = Pt(7)
tab_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
set_run_font(tab_run)

# PAGE field
for ftype in ('begin',):
    r = fp.add_run()
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), ftype)
    r._r.append(fc)

r = fp.add_run()
instr = OxmlElement('w:instrText')
instr.set(qn('xml:space'), 'preserve')
instr.text = 'PAGE'
r._r.append(instr)

for ftype in ('separate', 'end'):
    r = fp.add_run()
    fc = OxmlElement('w:fldChar')
    fc.set(qn('w:fldCharType'), ftype)
    r._r.append(fc)

# =============================================
#  TITLE — "Market Value Assessment" only
# =============================================

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
make_run(p_title, 'Market Value Assessment', size=18, bold=True, color=NAVY_TEXT)

doc.add_paragraph()  # spacer

# =============================================
#  GENERAL INFORMATION
# =============================================

section_heading(doc, '  GENERAL INFORMATION')

t = doc.add_table(rows=4, cols=6)

data = [
    ['Claim Number',    '{{MVR_CLAIM_NUMBER}}',  'Vehicle Owner',     '{{MVR_OWNER}}',            'Date',         '{{MVR_INSPECTION_DATE}}'],
    ['Carrier',         '{{MVR_CARRIER}}',       'Adjuster',          '{{MVR_ADJUSTER}}',         'Mileage',      '{{MVR_MILEAGE}}'],
    ['Year',            '{{MVR_YEAR}}',          'Make',              '{{MVR_MAKE}}',             'Model',        '{{MVR_MODEL}}'],
    ['VIN',             '',                      '',                  '',                         '',             ''],
]

for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        cell = t.rows[i].cells[j]
        is_label = (j % 2 == 0)
        if i == 3:
            continue  # handle VIN row separately
        if is_label and val:
            styled_cell(cell, val, bold=True, size=8, color=GRAY_BODY, fill=FILL_LABEL)
        else:
            styled_cell(cell, val, size=9)

# VIN row — label in col 0, token spans cols 1-5 via merge
vin_label = t.rows[3].cells[0]
styled_cell(vin_label, 'VIN', bold=True, size=8, color=GRAY_BODY, fill=FILL_LABEL)

# Merge cells 1 through 5 into one
merged = t.rows[3].cells[1]
for ci in range(2, 6):
    merged.merge(t.rows[3].cells[ci])
styled_cell(merged, '{{MVR_VIN}}', size=9)

doc.add_paragraph()

# =============================================
#  CONDITION
# =============================================

section_heading(doc, '  CONDITION')

ct = doc.add_table(rows=4, cols=2)
cond_rows = [
    ['Paint',             '{{MVR_COND_PAINT}}'],
    ['Glass',             '{{MVR_COND_GLASS}}'],
    ['Body / Sheet Metal','{{MVR_COND_BODY}}'],
    ['Interior',          '{{MVR_COND_INTERIOR}}'],
]
for i, (label, token) in enumerate(cond_rows):
    styled_cell(ct.rows[i].cells[0], label, bold=True, size=8, fill=FILL_LABEL)
    styled_cell(ct.rows[i].cells[1], token, size=9)

doc.add_paragraph()

# =============================================
#  DEALER COMPARABLES
# =============================================

section_heading(doc, '  DEALER COMPARABLES')

compt = doc.add_table(rows=8, cols=4)
comp_data = [
    ['',               'Comp 1',                   'Comp 2',                   'Comp 3'],
    ['Seller',         '{{MVR_COMP1_SELLER}}',      '{{MVR_COMP2_SELLER}}',      '{{MVR_COMP3_SELLER}}'],
    ['Location',       '{{MVR_COMP1_LOCATION}}',    '{{MVR_COMP2_LOCATION}}',    '{{MVR_COMP3_LOCATION}}'],
    ['Phone',          '{{MVR_COMP1_PHONE}}',       '{{MVR_COMP2_PHONE}}',       '{{MVR_COMP3_PHONE}}'],
    ['Asking Price',   '{{MVR_COMP1_PRICE}}',       '{{MVR_COMP2_PRICE}}',       '{{MVR_COMP3_PRICE}}'],
    ['Mileage Adj.',   '{{MVR_COMP1_MILEAGE_ADJ}}', '{{MVR_COMP2_MILEAGE_ADJ}}', '{{MVR_COMP3_MILEAGE_ADJ}}'],
    ['Condition Adj.', '{{MVR_COMP1_COND_ADJ}}',    '{{MVR_COMP2_COND_ADJ}}',    '{{MVR_COMP3_COND_ADJ}}'],
    ['Adjusted Price', '{{MVR_COMP1_ADJ_PRICE}}',   '{{MVR_COMP2_ADJ_PRICE}}',   '{{MVR_COMP3_ADJ_PRICE}}'],
]

for i, row_data in enumerate(comp_data):
    for j, val in enumerate(row_data):
        cell = compt.rows[i].cells[j]
        if i == 0:
            styled_cell(cell, val, bold=True, size=9, color=WHITE, fill=NAVY_FILL,
                        align=WD_ALIGN_PARAGRAPH.CENTER)
        elif i == 7:
            if j == 0:
                styled_cell(cell, val, bold=True, size=8, fill=FILL_HIGHLIGHT)
            else:
                styled_cell(cell, val, bold=True, size=9, fill=FILL_HIGHLIGHT,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
        elif j == 0:
            styled_cell(cell, val, bold=True, size=8, fill=FILL_LABEL)
        else:
            styled_cell(cell, val, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# =============================================
#  VALUATION SUMMARY
# =============================================

section_heading(doc, '  VALUATION SUMMARY')

vt = doc.add_table(rows=3, cols=4)
val_data = [
    ['Copart / Salvage Quote',        '{{MVR_SALVAGE_QUOTE}}',  'Unrelated Prior Damage', '{{MVR_PRIOR_DAMAGE}}'],
    ['Recommended ACV (avg of comps)','{{MVR_ACV}}',            'Local Tax Rate',         '{{MVR_TAX_RATE}}%'],
    ['Tax Amount',                    '{{MVR_TAX_AMOUNT}}',     'ACV + Tax (Total)',      '{{MVR_TOTAL}}'],
]

for i, row_data in enumerate(val_data):
    for j, val in enumerate(row_data):
        cell = vt.rows[i].cells[j]
        is_label = (j % 2 == 0)
        if i == 2:
            if is_label:
                styled_cell(cell, val, bold=True, size=8, fill=FILL_GOLD_TINT)
            else:
                styled_cell(cell, val, bold=True, size=9, fill=FILL_HIGHLIGHT)
        elif is_label:
            styled_cell(cell, val, bold=True, size=8, fill=FILL_LABEL)
        else:
            styled_cell(cell, val, size=9)

# Total value — larger, navy
total_run = vt.rows[2].cells[3].paragraphs[0].runs[0]
total_run.font.size = Pt(11)
total_run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

doc.add_paragraph()

# =============================================
#  ADDITIONAL REMARKS
# =============================================

section_heading(doc, '  ADDITIONAL REMARKS')

p = doc.add_paragraph()
make_run(p, '{{MVR_REMARKS}}', size=9)

doc.add_paragraph()

# =============================================
#  SIGNATURE BLOCK — horizontal rule + name + label + date
# =============================================

# Horizontal rule
hr = doc.add_paragraph()
add_para_border(hr, edge='bottom', color=BORDER_CELL, sz='2', space='4')
hr.paragraph_format.space_after = Pt(8)

# Appraiser name — italic, 16pt, navy
sig_para = doc.add_paragraph()
make_run(sig_para, '{{MVR_APPRAISER_NAME}}', size=16, italic=True, color=NAVY_TEXT)

# Label
sig_label = doc.add_paragraph()
make_run(sig_label, 'Inspector Signature', size=8, color=GRAY_SUB)

# Date
date_para = doc.add_paragraph()
make_run(date_para, 'Date: {{MVR_INSPECTION_DATE}}', size=9)

# =============================================
#  SAVE
# =============================================

out_dir = os.path.join('claim_cipher_app', 'forms', 'mvr')
os.makedirs(out_dir, exist_ok=True)
out_path = os.path.join(out_dir, 'MVR_TEMPLATE.docx')
doc.save(out_path)
print(f'MVR_TEMPLATE.docx created at {out_path}')

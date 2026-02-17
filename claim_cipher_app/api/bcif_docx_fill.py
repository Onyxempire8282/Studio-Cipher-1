#!/usr/bin/env python3
"""
BCIF DOCX filling endpoint for Total Loss V2.
"""

from io import BytesIO
from pathlib import Path
import re

from flask import request, jsonify, send_file
from docx import Document

from bcif_api import app

TEMPLATE_PATH = Path(__file__).parent.parent / "forms" / "bcif" / "BCIF_AUTOMATION_TEMPLATE_v3.docx"


def _replace_in_paragraph(paragraph, token_map):
    if not paragraph.text:
        return

    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join(run.text for run in runs)

    updated_text = full_text
    for key, value in token_map.items():
        token = "{{" + key + "}}"
        updated_text = updated_text.replace(token, value)

    updated_text = re.sub(r"\{\{.*?\}\}", "", updated_text)

    if updated_text == full_text:
        return

    runs[0].text = updated_text
    for run in runs[1:]:
        run.text = ""


def _replace_in_table(table, token_map):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, token_map)
            for nested in cell.tables:
                _replace_in_table(nested, token_map)


def _replace_in_container(container, token_map):
    for paragraph in container.paragraphs:
        _replace_in_paragraph(paragraph, token_map)
    for table in container.tables:
        _replace_in_table(table, token_map)


@app.route("/fill-bcif-docx", methods=["POST"])
def fill_bcif_docx():
    """
    Fill BCIF Word template using a JSON token map.
    """
    data = request.get_json(silent=True)
    if not data:
        return jsonify({"error": "Missing JSON payload"}), 400

    token_map = data.get("tokens", data)
    if not isinstance(token_map, dict):
        return jsonify({"error": "Token map must be an object"}), 400

    if not TEMPLATE_PATH.exists():
        return jsonify({"error": f"Template not found: {TEMPLATE_PATH}"}), 500

    normalized = {str(key): "" if value is None else str(value) for key, value in token_map.items()}

    doc = Document(TEMPLATE_PATH)

    _replace_in_container(doc, normalized)

    for section in doc.sections:
        _replace_in_container(section.header, normalized)
        _replace_in_container(section.footer, normalized)

    output_stream = BytesIO()
    doc.save(output_stream)
    output_stream.seek(0)

    return send_file(
        output_stream,
        as_attachment=True,
        download_name="BCIF_AUTOMATION_FILLED_v3.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
